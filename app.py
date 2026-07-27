"""
Mail Merge Automation - Backend
--------------------------------
Flask backend for:
  1. Parsing uploaded Excel/CSV data
  2. Parsing an uploaded Word (.docx) template and extracting {{tag}} placeholders
  3. Generating merged .docx documents (bulk, zipped)
  4. Sending merged documents as email attachments over SMTP
  5. Optionally saving a copy of every sent email into the mailbox's "Sent"
     folder via IMAP APPEND (SMTP alone does NOT do this automatically)
  6. Returning a per-recipient send log that the frontend can export as CSV

NOTE ON STATE: this app is designed to run on Vercel's serverless Python
runtime, which is stateless between requests (and may run on a different
instance/container each time). Nothing is stored on disk or in memory
across requests. The browser holds the parsed Excel data, the template
(base64) and the field mapping, and re-sends whatever a request needs.
Credentials (SMTP/IMAP) are supplied by the user in the browser each
session and are only ever held in server memory for the lifetime of a
single request - never written to disk or logged.
"""

import base64
import csv
import io
import os
import re
import smtplib
import ssl
import time
import zipfile
from datetime import datetime
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.utils import formataddr, make_msgid

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from flask import Flask, jsonify, request, send_file, send_from_directory

app = Flask(__name__, static_folder=None)

# Max upload size: 15 MB (Excel + docx templates are small; keeps abuse in check)
app.config["MAX_CONTENT_LENGTH"] = 15 * 1024 * 1024

TAG_PATTERN = re.compile(r"\{\{\s*([A-Za-z0-9_.\-]+)\s*\}\}")


# --------------------------------------------------------------------------
# Static frontend
# --------------------------------------------------------------------------
@app.route("/")
def index():
    return send_from_directory(os.path.dirname(os.path.abspath(__file__)), "index.html")


# --------------------------------------------------------------------------
# Helpers
# --------------------------------------------------------------------------
def infer_column_types(df: pd.DataFrame) -> dict:
    """Classify every column as 'number', 'date', or 'text'. This is what lets
    us later write real numbers/dates into the generated Excel and Word
    output instead of the plain-text strings that caused numbers and dates
    to show up left-aligned as text in the mail-merge attachments."""
    types = {}
    for col in df.columns:
        s = df[col]
        if pd.api.types.is_datetime64_any_dtype(s):
            types[col] = "date"
            continue
        if pd.api.types.is_numeric_dtype(s):
            types[col] = "number"
            continue

        non_null = s.dropna().astype(str).str.strip()
        non_null = non_null[non_null != ""]
        if len(non_null) == 0:
            types[col] = "text"
            continue

        as_num = pd.to_numeric(non_null, errors="coerce")
        if as_num.notna().mean() >= 0.9:
            types[col] = "number"
            continue

        as_date = pd.to_datetime(non_null, errors="coerce")
        if as_date.notna().mean() >= 0.9:
            types[col] = "date"
            continue

        types[col] = "text"
    return types


def format_cell_value(value, col_type: str = None) -> str:
    """Render a raw cell value for display in a generated Word doc/table.
    Prevents dates rendering as '2026-05-29 00:00:00' and numbers rendering
    as raw floats with long decimal tails."""
    if value is None or value == "":
        return ""
    if col_type == "date":
        try:
            ts = pd.to_datetime(value)
            if pd.isna(ts):
                return str(value)
            return ts.strftime("%d-%b-%Y")
        except Exception:
            return str(value)
    if col_type == "number":
        try:
            f = float(str(value).replace(",", "").strip())
            if f == int(f):
                return f"{int(f):,}"
            return f"{f:,.2f}"
        except (TypeError, ValueError):
            return str(value)
    return str(value)


def df_to_records(df: pd.DataFrame, column_types: dict = None):
    """Convert a dataframe to plain-JSON-safe list[dict].

    Unlike a blanket str(v) on every cell (which is what silently turned
    numbers and dates into plain text later on), this keeps number columns
    as real JSON numbers and date columns as bare 'YYYY-MM-DD' strings (no
    time-of-day) so the browser -> backend round trip preserves the type,
    and Excel/Word generation can format them properly instead of stamping
    them out as text.
    """
    column_types = column_types or {}
    records = []
    for _, row in df.iterrows():
        rec = {}
        for col, v in row.items():
            col = str(col)
            if pd.isna(v) or v == "":
                rec[col] = ""
                continue
            ctype = column_types.get(col)
            if ctype == "number":
                try:
                    f = float(v)
                    rec[col] = int(f) if float(f).is_integer() else f
                except (TypeError, ValueError):
                    rec[col] = str(v)
            elif ctype == "date":
                try:
                    ts = pd.to_datetime(v)
                    rec[col] = ts.strftime("%Y-%m-%d") if not pd.isna(ts) else str(v)
                except Exception:
                    rec[col] = str(v)
            else:
                rec[col] = str(v)
        records.append(rec)
    return records


def extract_tags_from_docx(doc: Document):
    """Find every {{tag}} used anywhere in the document body, tables, and headers/footers."""
    tags = set()

    def scan(text):
        for m in TAG_PATTERN.finditer(text):
            tags.add(m.group(1))

    for p in doc.paragraphs:
        scan(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    scan(p.text)
    for section in doc.sections:
        for part in (section.header, section.footer):
            for p in part.paragraphs:
                scan(p.text)

    return sorted(tags)


def replace_tags_in_paragraph(paragraph, mapping):
    """Replace {{tag}} occurrences in a paragraph while keeping the paragraph's
    formatting. Word frequently splits a single {{tag}} across multiple runs,
    so we rebuild the paragraph text as a whole and re-apply it to the first
    run (simple, reliable approach for mail-merge style templates)."""
    full_text = "".join(r.text for r in paragraph.runs)
    if "{{" not in full_text:
        return

    def sub(m):
        key = m.group(1)
        return str(mapping.get(key, m.group(0)))

    new_text = TAG_PATTERN.sub(sub, full_text)
    if new_text == full_text:
        return

    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.text = new_text


def fill_template(template_bytes: bytes, mapping: dict) -> bytes:
    """Return a new .docx (bytes) with every {{tag}} in the template replaced
    using `mapping` = {tag_name: value}."""
    doc = Document(io.BytesIO(template_bytes))

    for p in doc.paragraphs:
        replace_tags_in_paragraph(p, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_tags_in_paragraph(p, mapping)
    for section in doc.sections:
        for part in (section.header, section.footer):
            for p in part.paragraphs:
                replace_tags_in_paragraph(p, mapping)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def merge_text(template_str: str, row: dict, column_types: dict = None) -> str:
    """Substitute {{ColumnName}} directly against an excel row (used for
    email subject/body, which use the raw excel column names as tags)."""
    if not template_str:
        return ""
    column_types = column_types or {}

    def sub(m):
        key = m.group(1)
        if key not in row:
            return m.group(0)
        return format_cell_value(row.get(key), column_types.get(key))

    return TAG_PATTERN.sub(sub, template_str)


def safe_filename(name: str) -> str:
    name = re.sub(r"[^\w\-.]+", "_", str(name)).strip("_")
    return name or "document"


def tagify(name: str) -> str:
    """Turn an arbitrary column name into a valid {{tag}} identifier (only
    letters/digits/underscore/dot/dash, per TAG_PATTERN). Column names in
    real spreadsheets very often contain spaces or punctuation (e.g. 'Bill
    Amt'), which can't appear inside a literal {{tag}} — this lets aggregate
    tags like {{Total_Bill_Amt}} still work off of such columns."""
    safe = re.sub(r"[^A-Za-z0-9_.\-]+", "_", str(name)).strip("_")
    return safe or "Field"


def build_case_zip(template_bytes: bytes, rows: list, doc_mapping: dict, filename_field: str, default_prefix: str = "case", column_types: dict = None) -> bytes:
    """Merge the Word template against every row and bundle the results into
    a single ZIP (used both by /api/generate and the company-wise 'zipped
    individual case documents' attachment)."""
    column_types = column_types or {}
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        used_names = set()
        for i, row in enumerate(rows, start=1):
            row_mapping = {
                tag: format_cell_value(row.get(col, ""), column_types.get(col))
                for tag, col in doc_mapping.items()
            }
            doc_bytes = fill_template(template_bytes, row_mapping)

            base_name = safe_filename(row.get(filename_field, "")) if filename_field else f"{default_prefix}_{i}"
            fname = f"{base_name}.docx"
            n = 1
            while fname in used_names:
                n += 1
                fname = f"{base_name}_{n}.docx"
            used_names.add(fname)

            zf.writestr(fname, doc_bytes)

    zip_buffer.seek(0)
    return zip_buffer.getvalue()


def build_filtered_excel(rows: list, columns: list, column_types: dict = None) -> bytes:
    """Build a .xlsx containing only the given rows/columns (a single
    company's cases) — the 'Filtered Excel' company-wise attachment.

    Number and date columns are coerced back to real dtypes and given an
    explicit Excel number format, instead of being written as plain text
    (which is what made every number/date show up left-aligned as text in
    the mailed attachment)."""
    column_types = column_types or {}
    df = pd.DataFrame(rows)
    cols = [c for c in columns if c in df.columns] or list(df.columns)
    df = df[cols]

    for col in cols:
        ctype = column_types.get(col)
        if ctype == "number":
            df[col] = pd.to_numeric(
                df[col].astype(str).str.replace(",", "", regex=False).replace("", pd.NA),
                errors="coerce",
            )
        elif ctype == "date":
            df[col] = pd.to_datetime(df[col].replace("", pd.NaT), errors="coerce")

    out = io.BytesIO()
    with pd.ExcelWriter(out, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Cases")
        ws = writer.sheets["Cases"]
        for idx, col in enumerate(cols, start=1):
            ctype = column_types.get(col)
            if ctype not in ("number", "date"):
                continue
            number_format = "#,##0.00" if ctype == "number" else "dd-mmm-yyyy"
            for row_cells in ws.iter_rows(min_row=2, min_col=idx, max_col=idx):
                for cell in row_cells:
                    cell.number_format = number_format
    return out.getvalue()


def _shade_cell(cell, hex_color: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _repeat_header_row(row):
    """Mark a table row so it repeats as a header on every page."""
    trPr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trPr.append(header)


def build_table_docx(group_label: str, rows: list, columns: list, column_types: dict = None) -> bytes:
    """Build a .docx with a single aligned, bordered table listing every case
    for one company — the 'Summary Table Document' company-wise attachment.

    Real-world case files often have 15-25+ columns; left at Word's default
    "autofit to window" behaviour in a portrait page, that many columns get
    squeezed unreadably narrow and the table looks broken. For wide tables
    this switches the page to landscape with slim margins, fixes the table
    layout, and gives every column an explicit equal share of the page width
    (plus a smaller font) so it actually renders as an aligned table."""
    column_types = column_types or {}
    doc = Document()

    cols = columns or (list(rows[0].keys()) if rows else [])
    n_cols = max(len(cols), 1)
    wide = n_cols > 6

    section = doc.sections[0]
    if wide:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)

    doc.add_heading(f"Case Summary — {group_label}", level=1)
    sub = doc.add_paragraph(f"Total cases: {len(rows)}")
    if sub.runs:
        sub.runs[0].italic = True

    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    table.autofit = False

    # Force a fixed column layout — Word otherwise ignores explicit column
    # widths on tables and re-autofits them to content, which is what broke
    # the alignment on wide tables.
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    usable_width = section.page_width - section.left_margin - section.right_margin
    col_width = int(usable_width / n_cols)
    font_size = Pt(9) if n_cols <= 10 else (Pt(8) if n_cols <= 16 else Pt(7))

    for col_obj in table.columns:
        col_obj.width = col_width

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(cols):
        hdr_cells[i].width = col_width
        hdr_cells[i].text = str(col)
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = font_size
        _shade_cell(hdr_cells[i], "D9E2F3")
    _repeat_header_row(table.rows[0])

    for row in rows:
        cells = table.add_row().cells
        for i, col in enumerate(cols):
            cells[i].width = col_width
            cells[i].text = format_cell_value(row.get(col, ""), column_types.get(col))
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = font_size

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def sum_column(rows: list, col: str):
    """Sum a column across a list of rows. Returns (total, ok) where ok is
    False if any non-blank cell in that column couldn't be parsed as a
    number (mirrors the parsing rules already used for amount_field)."""
    total = 0.0
    ok = True
    for row in rows:
        raw = str(row.get(col, "")).replace(",", "").strip()
        if not raw:
            continue
        try:
            total += float(raw)
        except ValueError:
            ok = False
    return total, ok


def resolve_letter_tags(template_tags, tag_aggregates: dict, ctx: dict, rows: list, mapping: dict) -> dict:
    """Build the final {tag: value} mapping used to fill a company-wise
    letter template. For each tag found in the template:
      1. If the user explicitly told us how to aggregate this tag
         (tag_aggregates[tag] — count / sum of a column / difference of two
         column sums), compute that directly. This is what lets a template
         use ANY tag name (e.g. {{Total_Claimed_Amount}}) for a real sum,
         instead of requiring the tag to literally be named
         {{Total_<ExcelColumnName>}}.
      2. Otherwise, if it's one of the automatic aggregate tags
         (Case_Count, Group_Label, a group-by column, {{Total_<amount_field>}}...)
         use that.
      3. Otherwise fall back to the plain per-tag column mapping (the first
         case's raw value) — for static info like company address/email.
    """
    tag_aggregates = tag_aggregates or {}
    row_mapping = {}
    for tag in template_tags:
        agg = tag_aggregates.get(tag)
        if agg:
            kind = agg.get("type")
            if kind == "count":
                row_mapping[tag] = str(len(rows))
            elif kind == "sum":
                total, ok = sum_column(rows, agg.get("column", ""))
                row_mapping[tag] = f"{total:,.2f}" if ok else "N/A"
            elif kind == "diff":
                a_total, a_ok = sum_column(rows, agg.get("column_a", ""))
                b_total, b_ok = sum_column(rows, agg.get("column_b", ""))
                row_mapping[tag] = f"{(a_total - b_total):,.2f}" if (a_ok and b_ok) else "N/A"
            else:
                row_mapping[tag] = ""
            continue

        if tag in ctx:
            row_mapping[tag] = ctx[tag]
        else:
            col = mapping.get(tag)
            row_mapping[tag] = ctx.get(col, "") if col else ""
    return row_mapping


def compute_group_aggregates(group_fields, group_key_values: dict, rows: list, recipient_email: str, amount_field: str = None, column_types: dict = None) -> dict:
    """Build the merge context available to the template/subject/body when
    sending or generating one document per group: every plain excel column
    (taken from the group's FIRST row, formatted per column_types, since a
    single document can't hold N different values), plus the aggregate tags
    {{<each group_field>}}, {{Case_Count}}, {{Company_Email}}, {{Group_Label}}
    and {{Total_<amount_field>}} if an amount column was chosen. Aggregate
    tags always take priority over a same-named plain column.

    `group_fields` may be a single column name (legacy, single-field grouping
    used by the Email Dispatch tab) or a list of column names (multi-field
    grouping, e.g. Company + Aging bucket, used by Document Generation)."""
    column_types = column_types or {}
    if isinstance(group_fields, str):
        group_fields = [group_fields] if group_fields else []

    ctx = {}
    if rows:
        for k, v in rows[0].items():
            ctx[str(k)] = format_cell_value(v, column_types.get(str(k)))

    for f in group_fields:
        val = group_key_values.get(f, "")
        ctx[f] = val
        ctx[tagify(f)] = val
    ctx["Group_Label"] = " | ".join(str(group_key_values.get(f, "")) for f in group_fields)
    ctx["Case_Count"] = str(len(rows))
    ctx["Company_Email"] = recipient_email

    if amount_field:
        total = 0.0
        ok = True
        for row in rows:
            raw = str(row.get(amount_field, "")).replace(",", "").strip()
            if not raw:
                continue
            try:
                total += float(raw)
            except ValueError:
                ok = False
        total_str = f"{total:,.2f}" if ok else "N/A"
        ctx[f"Total_{amount_field}"] = total_str
        ctx[f"Total_{tagify(amount_field)}"] = total_str

    return ctx


def open_smtp(conf):
    host = conf["host"]
    port = int(conf["port"])
    security = conf.get("security", "starttls")  # 'ssl' | 'starttls' | 'none'

    if security == "ssl":
        server = smtplib.SMTP_SSL(host, port, timeout=20, context=ssl.create_default_context())
    else:
        server = smtplib.SMTP(host, port, timeout=20)
        server.ehlo()
        if security == "starttls":
            server.starttls(context=ssl.create_default_context())
            server.ehlo()

    if conf.get("username"):
        server.login(conf["username"], conf.get("password", ""))
    return server


def open_imap(conf):
    import imaplib

    host = conf["host"]
    port = int(conf.get("port") or 993)
    if conf.get("security", "ssl") == "ssl":
        m = imaplib.IMAP4_SSL(host, port, timeout=20)
    else:
        m = imaplib.IMAP4(host, port, timeout=20)
        if conf.get("security") == "starttls":
            m.starttls(ssl_context=ssl.create_default_context())
    m.login(conf["username"], conf.get("password", ""))
    return m


def append_to_sent(imap_conf, raw_message: bytes):
    import imaplib

    m = open_imap(imap_conf)
    try:
        folder = imap_conf.get("folder") or "Sent"
        typ, _ = m.select(folder, readonly=False)
        if typ != "OK":
            # try common alternates if the given folder name doesn't exist
            for alt in ("Sent", "INBOX.Sent", "Sent Items", "INBOX.Sent Items"):
                typ, _ = m.select(alt, readonly=False)
                if typ == "OK":
                    break
        m.append(folder, "\\Seen", imaplib.Time2Internaldate(time.time()), raw_message)
    finally:
        try:
            m.logout()
        except Exception:
            pass


def build_email(from_name, from_email, to_email, subject, html_body, attachment=None, attachment_name=None, attachments=None):
    """Build a MIME email. `attachment`/`attachment_name` is the original
    single-attachment path (per-case mode). `attachments` is an optional list
    of (bytes, filename) tuples for callers that need more than one file
    attached (company-wise mode: table doc + excel + zip together)."""
    msg = MIMEMultipart("mixed")
    msg["Subject"] = subject
    msg["From"] = formataddr((from_name or from_email, from_email))
    msg["To"] = to_email
    msg["Message-ID"] = make_msgid()

    alt = MIMEMultipart("alternative")
    plain = re.sub("<[^<]+?>", "", html_body or "")
    alt.attach(MIMEText(plain, "plain", "utf-8"))
    alt.attach(MIMEText(html_body or "", "html", "utf-8"))
    msg.attach(alt)

    all_attachments = list(attachments or [])
    if attachment is not None:
        all_attachments.append((attachment, attachment_name))

    for att_bytes, att_name in all_attachments:
        part = MIMEApplication(att_bytes, Name=att_name)
        part["Content-Disposition"] = f'attachment; filename="{att_name}"'
        msg.attach(part)

    return msg


# --------------------------------------------------------------------------
# API: upload spreadsheet
# --------------------------------------------------------------------------
@app.route("/api/upload-data", methods=["POST"])
def upload_data():
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "No file uploaded"}), 400

    filename = f.filename or ""
    try:
        if filename.lower().endswith(".csv"):
            df = pd.read_csv(f)
        else:
            df = pd.read_excel(f)
    except Exception as e:
        return jsonify({"error": f"Could not read spreadsheet: {e}"}), 400

    if df.empty:
        return jsonify({"error": "Spreadsheet has no rows"}), 400

    df.columns = [str(c).strip() for c in df.columns]
    column_types = infer_column_types(df)
    records = df_to_records(df, column_types)

    return jsonify({
        "filename": filename,
        "columns": list(df.columns),
        "column_types": column_types,
        "rows": records,
        "row_count": len(records),
    })


# --------------------------------------------------------------------------
# API: upload word template
# --------------------------------------------------------------------------
@app.route("/api/upload-template", methods=["POST"])
def upload_template():
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "No file uploaded"}), 400
    if not (f.filename or "").lower().endswith(".docx"):
        return jsonify({"error": "Please upload a .docx file (old .doc format is not supported)"}), 400

    file_bytes = f.read()
    try:
        doc = Document(io.BytesIO(file_bytes))
        tags = extract_tags_from_docx(doc)
    except Exception as e:
        return jsonify({"error": f"Could not read Word template: {e}"}), 400

    if not tags:
        return jsonify({"error": "No {{tags}} found in this template. Add placeholders like {{Full_Name}} and re-upload."}), 400

    return jsonify({
        "filename": f.filename,
        "tags": tags,
        "template_b64": base64.b64encode(file_bytes).decode("ascii"),
    })


# --------------------------------------------------------------------------
# API: generate merged documents (zip)
#
# mode = "per-case" (default, unchanged behaviour): one .docx per row.
#
# mode = "company": group ALL rows by one or more columns (e.g. Insurance
# Company, or Company + Aging bucket together) and produce ONE set of
# documents per group instead of one per case — so 1000 cases across 10
# companies produces 10 (not 1000) documents. Per group you can request any
# combination of:
#   - a "letter" .docx: the uploaded Word template filled with aggregate
#     tags ({{Case_Count}}, {{Total_<amount_field>}}, the group-by column(s))
#   - a Summary Table Document .docx (one row per case, one doc per group)
#   - a filtered Excel .xlsx (just that group's rows)
#   - a .zip of the individual per-case .docx documents for that group
# All group outputs are bundled together into a single downloaded .zip.
# --------------------------------------------------------------------------
@app.route("/api/generate", methods=["POST"])
def generate_documents():
    data = request.get_json(force=True, silent=True) or {}
    template_b64 = data.get("template_b64")
    mapping = data.get("mapping") or {}       # {tag: excel_column}
    rows = data.get("rows") or []
    name_field = data.get("filename_field")   # excel column to use for output filenames
    column_types = data.get("column_types") or {}
    mode = data.get("mode") or "per-case"

    if not rows:
        return jsonify({"error": "No data rows to generate from"}), 400

    if mode != "company":
        if not template_b64:
            return jsonify({"error": "Missing template"}), 400
        try:
            template_bytes = base64.b64decode(template_b64)
        except Exception:
            return jsonify({"error": "Invalid template data"}), 400

        zip_bytes = build_case_zip(template_bytes, rows, mapping, name_field, default_prefix="document", column_types=column_types)

        return send_file(
            io.BytesIO(zip_bytes),
            mimetype="application/zip",
            as_attachment=True,
            download_name=f"merged_documents_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
        )

    # ---- Company-wise / multi-column group-wise generation ----
    group_fields = data.get("group_fields") or []
    if not group_fields:
        return jsonify({"error": "Pick at least one column to group by"}), 400
    amount_field = data.get("amount_field") or None
    recipient_field = data.get("recipient_field") or None  # optional, only feeds {{Company_Email}}
    tag_aggregates = data.get("tag_aggregates") or {}  # {{tag}}: {"type": "count"|"sum"|"diff", ...}

    attach_letter = bool(data.get("attach_letter"))
    attach_table = bool(data.get("attach_table"))
    attach_excel = bool(data.get("attach_excel"))
    attach_zip = bool(data.get("attach_zip"))

    if not any([attach_letter, attach_table, attach_excel, attach_zip]):
        return jsonify({"error": "Pick at least one document type to generate per group"}), 400

    table_columns = data.get("table_columns") or list(rows[0].keys())
    excel_columns = data.get("excel_columns") or list(rows[0].keys())

    template_bytes = None
    if attach_letter or attach_zip:
        if not template_b64:
            return jsonify({"error": "Upload a Word template first (needed for the letter / zipped case documents), or turn those off"}), 400
        try:
            template_bytes = base64.b64decode(template_b64)
        except Exception:
            return jsonify({"error": "Invalid template data"}), 400

    groups = {}
    order = []
    for row in rows:
        key_values = {f: (str(row.get(f, "")).strip() or "(blank)") for f in group_fields}
        key = " | ".join(key_values[f] for f in group_fields)
        if key not in groups:
            groups[key] = {"key_values": key_values, "rows": []}
            order.append(key)
        groups[key]["rows"].append(row)

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for key in order:
            g = groups[key]
            g_rows = g["rows"]
            label = safe_filename(key)

            recipient_email = ""
            if recipient_field:
                emails = [str(r.get(recipient_field, "")).strip() for r in g_rows if str(r.get(recipient_field, "")).strip()]
                recipient_email = emails[0] if emails else ""

            if attach_letter:
                ctx = compute_group_aggregates(group_fields, g["key_values"], g_rows, recipient_email, amount_field, column_types)
                template_tags = extract_tags_from_docx(Document(io.BytesIO(template_bytes)))
                row_mapping = resolve_letter_tags(template_tags, tag_aggregates, ctx, g_rows, mapping)
                doc_bytes = fill_template(template_bytes, row_mapping)
                zf.writestr(f"{label}_letter.docx", doc_bytes)

            if attach_table:
                table_bytes = build_table_docx(key, g_rows, table_columns, column_types)
                zf.writestr(f"{label}_summary.docx", table_bytes)

            if attach_excel:
                excel_bytes = build_filtered_excel(g_rows, excel_columns, column_types)
                zf.writestr(f"{label}_cases.xlsx", excel_bytes)

            if attach_zip:
                case_zip_bytes = build_case_zip(template_bytes, g_rows, mapping, name_field, default_prefix=label, column_types=column_types)
                zf.writestr(f"{label}_case_documents.zip", case_zip_bytes)

    zip_buffer.seek(0)
    return send_file(
        zip_buffer,
        mimetype="application/zip",
        as_attachment=True,
        download_name=f"group_documents_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
    )


# --------------------------------------------------------------------------
# API: test SMTP / IMAP credentials
# --------------------------------------------------------------------------
@app.route("/api/test-connection", methods=["POST"])
def test_connection():
    data = request.get_json(force=True, silent=True) or {}
    kind = data.get("type")  # 'smtp' | 'imap'
    conf = data.get("config") or {}

    try:
        if kind == "smtp":
            server = open_smtp(conf)
            server.quit()
        elif kind == "imap":
            m = open_imap(conf)
            m.logout()
        else:
            return jsonify({"ok": False, "error": "Unknown connection type"}), 400
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 200

    return jsonify({"ok": True})


# --------------------------------------------------------------------------
# API: send emails (called by the frontend in small batches)
# --------------------------------------------------------------------------
@app.route("/api/send-emails", methods=["POST"])
def send_emails():
    data = request.get_json(force=True, silent=True) or {}

    smtp_conf = data.get("smtp") or {}
    from_name = data.get("from_name") or ""
    from_email = data.get("from_email") or smtp_conf.get("username", "")
    subject_template = data.get("subject") or ""
    body_template = data.get("body") or ""
    recipient_field = data.get("recipient_field")
    rows = data.get("rows") or []
    column_types = data.get("column_types") or {}

    save_copy = bool(data.get("save_sent_copy"))
    imap_conf = data.get("imap") or {}

    attach_doc = bool(data.get("attach_document"))
    template_b64 = data.get("template_b64")
    doc_mapping = data.get("doc_mapping") or {}
    filename_field = data.get("filename_field")

    if not recipient_field:
        return jsonify({"error": "recipient_field (which column holds the email address) is required"}), 400
    if not rows:
        return jsonify({"error": "No rows to send"}), 400

    template_bytes = None
    if attach_doc and template_b64:
        try:
            template_bytes = base64.b64decode(template_b64)
        except Exception:
            return jsonify({"error": "Invalid template data"}), 400

    log = []
    try:
        smtp_server = open_smtp(smtp_conf)
    except Exception as e:
        return jsonify({"error": f"Could not connect to SMTP server: {e}"}), 400

    for row in rows:
        to_email = str(row.get(recipient_field, "")).strip()
        entry = {
            "recipient": to_email,
            "subject": "",
            "status": "failed",
            "error": "",
            "timestamp": datetime.now().isoformat(timespec="seconds"),
        }

        if not to_email or "@" not in to_email:
            entry["error"] = "Missing/invalid email address"
            log.append(entry)
            continue

        try:
            subject = merge_text(subject_template, row, column_types)
            body = merge_text(body_template, row, column_types)
            entry["subject"] = subject

            attachment_bytes = None
            attachment_name = None
            if attach_doc and template_bytes is not None:
                row_mapping = {
                    tag: format_cell_value(row.get(col, ""), column_types.get(col))
                    for tag, col in doc_mapping.items()
                }
                attachment_bytes = fill_template(template_bytes, row_mapping)
                base_name = safe_filename(row.get(filename_field, "")) if filename_field else safe_filename(to_email)
                attachment_name = f"{base_name}.docx"

            msg = build_email(from_name, from_email, to_email, subject, body, attachment_bytes, attachment_name)
            raw = msg.as_bytes()

            smtp_server.sendmail(from_email, [to_email], raw)
            entry["status"] = "sent"

            if save_copy:
                try:
                    append_to_sent(imap_conf, raw)
                    entry["sent_copy_saved"] = True
                except Exception as e:
                    entry["sent_copy_saved"] = False
                    entry["sent_copy_error"] = str(e)

        except Exception as e:
            entry["status"] = "failed"
            entry["error"] = str(e)

        log.append(entry)

    try:
        smtp_server.quit()
    except Exception:
        pass

    sent = sum(1 for e in log if e["status"] == "sent")
    failed = len(log) - sent
    return jsonify({"log": log, "sent": sent, "failed": failed})


# --------------------------------------------------------------------------
# API: send ONE company-wise email (all of that company's cases grouped
# together). Insurance-recovery use case: instead of one email per patient,
# HDFC / ICICI Lombard / Care / Niva Bupa etc. each get a single email
# carrying only their own cases, with up to 3 attachments:
#   1. Summary Table Document (.docx) - one aligned table of every case
#   2. Filtered Excel (.xlsx)          - only that company's rows
#   3. Zipped individual case docs     - the existing Word-template merge,
#      one .docx per case, bundled into a .zip (optional)
# The frontend calls this once per company (like /api/send-emails is called
# once per small batch) so the Email Shot Status panel can update live.
# --------------------------------------------------------------------------
@app.route("/api/send-company-group", methods=["POST"])
def send_company_group():
    data = request.get_json(force=True, silent=True) or {}

    smtp_conf = data.get("smtp") or {}
    from_name = data.get("from_name") or ""
    from_email = data.get("from_email") or smtp_conf.get("username", "")
    subject_template = data.get("subject") or ""
    body_template = data.get("body") or ""

    group_field = data.get("group_field") or ""
    group_value = data.get("group_value", "")
    rows = data.get("rows") or []
    recipient_email = str(data.get("recipient_email") or "").strip()
    amount_field = data.get("amount_field") or None
    column_types = data.get("column_types") or {}

    save_copy = bool(data.get("save_sent_copy"))
    imap_conf = data.get("imap") or {}

    attach_table = bool(data.get("attach_table"))
    table_columns = data.get("table_columns") or (list(rows[0].keys()) if rows else [])

    attach_excel = bool(data.get("attach_excel"))
    excel_columns = data.get("excel_columns") or (list(rows[0].keys()) if rows else [])

    attach_zip = bool(data.get("attach_zip"))
    template_b64 = data.get("template_b64")
    doc_mapping = data.get("doc_mapping") or {}
    filename_field = data.get("filename_field")

    entry = {
        "group": group_value,
        "case_count": len(rows),
        "recipient": recipient_email,
        "subject": "",
        "status": "failed",
        "error": "",
        "timestamp": datetime.now().isoformat(timespec="seconds"),
    }

    if not rows:
        entry["error"] = "No cases found in this company group"
        return jsonify({"log": [entry]})
    if not recipient_email or "@" not in recipient_email:
        entry["error"] = "Missing/invalid recipient email for this company"
        return jsonify({"log": [entry]})

    try:
        ctx = compute_group_aggregates([group_field] if group_field else [], {group_field: group_value} if group_field else {}, rows, recipient_email, amount_field, column_types)
        subject = merge_text(subject_template, ctx)
        body = merge_text(body_template, ctx)
        entry["subject"] = subject

        attachments = []

        if attach_table:
            table_bytes = build_table_docx(str(group_value), rows, table_columns, column_types)
            attachments.append((table_bytes, f"{safe_filename(group_value)}_summary.docx"))

        if attach_excel:
            excel_bytes = build_filtered_excel(rows, excel_columns, column_types)
            attachments.append((excel_bytes, f"{safe_filename(group_value)}_cases.xlsx"))

        if attach_zip and template_b64:
            template_bytes = base64.b64decode(template_b64)
            zip_bytes = build_case_zip(template_bytes, rows, doc_mapping, filename_field, column_types=column_types)
            attachments.append((zip_bytes, f"{safe_filename(group_value)}_case_documents.zip"))

        msg = build_email(from_name, from_email, recipient_email, subject, body, attachments=attachments)
        raw = msg.as_bytes()

        try:
            smtp_server = open_smtp(smtp_conf)
        except Exception as e:
            entry["error"] = f"Could not connect to SMTP server: {e}"
            return jsonify({"log": [entry]})

        try:
            smtp_server.sendmail(from_email, [recipient_email], raw)
            entry["status"] = "sent"
        finally:
            try:
                smtp_server.quit()
            except Exception:
                pass

        if entry["status"] == "sent" and save_copy:
            try:
                append_to_sent(imap_conf, raw)
                entry["sent_copy_saved"] = True
            except Exception as e:
                entry["sent_copy_saved"] = False
                entry["sent_copy_error"] = str(e)

    except Exception as e:
        entry["status"] = "failed"
        entry["error"] = str(e)

    return jsonify({"log": [entry]})


# --------------------------------------------------------------------------
# API: log download as CSV (kept server-side too, in case the frontend
# wants to re-download a log it already has instead of rebuilding client-side)
# --------------------------------------------------------------------------
@app.route("/api/download-log", methods=["POST"])
def download_log():
    data = request.get_json(force=True, silent=True) or {}
    log = data.get("log") or []

    buf = io.StringIO()
    fieldnames = ["timestamp", "recipient", "subject", "status", "sent_copy_saved", "error"]
    writer = csv.DictWriter(buf, fieldnames=fieldnames, extrasaction="ignore")
    writer.writeheader()
    for row in log:
        writer.writerow(row)

    mem = io.BytesIO(buf.getvalue().encode("utf-8-sig"))
    return send_file(
        mem,
        mimetype="text/csv",
        as_attachment=True,
        download_name=f"send_log_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
    )


if __name__ == "__main__":
    app.run(debug=True, port=5000)
